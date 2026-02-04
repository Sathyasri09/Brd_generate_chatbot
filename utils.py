import os
from htmldocx import HtmlToDocx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def convert_html_to_docx(html_content: str, output_filename: str = "generated_brd.docx") -> str:
    """
    Standard Document Engine:
    Handles centering of title page and inserts page breaks where marked.
    Saves to external directory to prevent Chainlit watch/reload loop.
    """
    new_parser = HtmlToDocx()
    doc = Document()
    
    if not html_content:
        return ""
    
    # Standard marker for splits
    marker = "<div class='page-break'></div>"
    
    if marker in html_content:
        parts = html_content.split(marker)
        for i, part in enumerate(parts):
            if not part.strip():
                continue
            
            # Add this content block
            new_parser.add_html_to_document(part, doc)
            
            # Add physical Word page break
            if i < len(parts) - 1:
                doc.add_page_break()
    else:
        new_parser.add_html_to_document(html_content, doc)
    
    # Executive Styling Pass
    title_keywords = ["BUSINESS REQUIREMENT DOCUMENT", "TECHNICAL SPECIFICATION DOCUMENT", "VERSION 1.0", "DATE:", "PROJECT NAME:"]
    
    for para in doc.paragraphs[:25]:
        text_upper = para.text.upper()
        
        # Center Title Page Info (First page blocks)
        if any(k in text_upper for k in title_keywords):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(12)
            for run in para.runs:
                if "DOCUMENT" in text_upper:
                    run.font.size = Pt(28)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(26, 54, 93)
                else:
                    run.font.size = Pt(14)
                    run.font.bold = True
        
        # Center "Table of Contents" Header (if present)
        if "TABLE OF CONTENTS" in text_upper:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(20)
            para.paragraph_format.space_after = Pt(20)
            for run in para.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 54, 93)

    # Global styles for professional consistency
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            para.keep_with_next = True
            for run in para.runs:
                run.font.color.rgb = RGBColor(26, 54, 93)

    # Highlight all tables
    for table in doc.tables:
        table.style = 'Table Grid'
    
    # Save to directory OUTSIDE the project root to prevent Chainlit watch/reload loop
    output_dir = os.path.abspath(os.path.join(os.getcwd(), "..", "brd_outputs"))
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    return output_path

def load_template(template_path: str) -> str:
    """Loads the BRD text template."""
    with open(template_path, 'r', encoding='utf-8') as f:
        return f.read()

def parse_brd_sections(template_content: str) -> list[str]:
    """Identifies key sections for the BA prompt."""
    return [
        "1 Business Need", "2 Business Requirements", "3 Key Success Factors",
        "4 Assumptions, Dependencies and Constraints", "5 User Acceptance Criteria",
        "6 Risks and Impact", "7 References", "8 Solution Overview"
    ]
